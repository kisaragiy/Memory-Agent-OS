# -*- coding: utf-8 -*-
"""Generate architecture PNGs and portfolio docx under docs/."""
from __future__ import annotations

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DIAGRAMS = DOCS / "diagrams"

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    print("Missing dependency: pip install python-docx")
    sys.exit(1)

try:
    from PIL import Image, ImageDraw, ImageFont
except ImportError:
    print("Missing dependency: pip install Pillow")
    sys.exit(1)


def _font(size: int = 16):
    candidates = [
        "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
    ]
    for p in candidates:
        if Path(p).is_file():
            return ImageFont.truetype(p, size)
    return ImageFont.load_default()


def draw_layers_diagram(out: Path) -> None:
    """Six-layer stack diagram."""
    w, h = 900, 720
    img = Image.new("RGB", (w, h), "#f8f9fb")
    draw = ImageDraw.Draw(img)
    title_f = _font(22)
    box_f = _font(15)
    small_f = _font(12)
    draw.text((w // 2 - 180, 24), "Memory Agent OS — 六层架构", fill="#1a1a2e", font=title_f)
    layers = [
        ("Presentation", "React Web UI · FastAPI :8787", "#4a90d9"),
        ("Control", "ModelPolicy · MemoryControl · user/dev/debug", "#5b8def"),
        ("Orchestration", "KernelOrchestrator · AutonomousOSLoop", "#6c63ff"),
        ("Execution", "ExecutionEngine → ToolRouter → Tools", "#7b61c4"),
        ("State", "MemoryLayer · Chroma/Ollama 可选", "#9b59b6"),
        ("World", "WorldRuntime · NarrativeWriter", "#8e44ad"),
    ]
    x0, bw, bh, gap = 120, 660, 72, 18
    y = 80
    for name, desc, color in layers:
        draw.rounded_rectangle([x0, y, x0 + bw, y + bh], radius=10, fill=color, outline="#333")
        draw.text((x0 + 20, y + 12), name, fill="white", font=box_f)
        draw.text((x0 + 20, y + 40), desc, fill="#eef", font=small_f)
        y += bh + gap
        if y < h - 60:
            cx = x0 + bw // 2
            draw.polygon([(cx - 8, y - gap + 4), (cx + 8, y - gap + 4), (cx, y - 2)], fill="#555")
    draw.text((x0, h - 44), "约束：唯一 ExecutionEngine · 唯一 MemoryLayer 入口", fill="#444", font=small_f)
    out.parent.mkdir(parents=True, exist_ok=True)
    img.save(out, "PNG")


def draw_execution_path(out: Path) -> None:
    """Horizontal pipeline for entry()."""
    w, h = 1000, 380
    img = Image.new("RGB", (w, h), "#ffffff")
    draw = ImageDraw.Draw(img)
    f = _font(13)
    t = _font(18)
    draw.text((24, 16), "唯一执行路径 main.py → AgentOSRuntime.entry()", fill="#1a1a2e", font=t)
    steps = [
        "InputNormalizer",
        "MemoryLayer\nbuild_context",
        "WorldRuntime\nstep",
        "KernelOrchestrator\nbuild_plan",
        "ExecutionEngine",
        "ToolRouter",
        "record_episode",
    ]
    sw, sh, gap = 118, 56, 12
    x, y = 24, 70
    for i, label in enumerate(steps):
        draw.rounded_rectangle([x, y, x + sw, y + sh], radius=6, fill="#e8f4fc", outline="#2c6fad")
        for j, line in enumerate(label.split("\n")):
            draw.text((x + 8, y + 10 + j * 18), line, fill="#1a1a2e", font=f)
        if i < len(steps) - 1:
            ax = x + sw + 2
            draw.line([(ax, y + sh // 2), (ax + gap - 2, y + sh // 2)], fill="#2c6fad", width=2)
            draw.polygon([(ax + gap, y + sh // 2), (ax + gap - 8, y + sh // 2 - 5), (ax + gap - 8, y + sh // 2 + 5)], fill="#2c6fad")
        x += sw + gap
    draw.text((24, 150), "Phase5: AutonomousOSLoop 每步仅调用 entry() — 禁止第二 ExecutionEngine", fill="#555", font=f)
    draw.text((24, 180), "记忆治理: UI/API → MemoryControl → execute_memory_op → apply_mutation → trace", fill="#555", font=f)
    out.parent.mkdir(parents=True, exist_ok=True)
    img.save(out, "PNG")


def _heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def _bullet(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _add_image(doc: Document, path: Path, caption: str, width_in: float = 6.2) -> None:
    if path.is_file():
        doc.add_picture(str(path), width=Inches(width_in))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cap.runs:
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def build_architecture_docx(layers: Path, pipeline: Path) -> Path:
    out = DOCS / "MemoryAgentOS-系统架构说明.docx"
    doc = Document()
    _heading(doc, "Memory Agent OS — 系统架构说明", 0)
    _para(doc, "版本：2026.05 · 单内核 Agent Runtime + 虚拟叙事世界引擎")
    doc.add_paragraph()

    _heading(doc, "1. 项目定位", 1)
    _bullet(
        doc,
        [
            "LLM 驱动的操作系统级 Agent Runtime，非 Chatbot / Workflow 封装",
            "唯一执行路径：main.py → AgentOSRuntime.entry() → ExecutionEngine",
            "治理化记忆：显式写入经 execute_memory_op，全链路 trace",
            "Phase 5：observe → plan → act → reflect 自主闭环",
        ],
    )

    _heading(doc, "2. 六层工程架构", 1)
    _add_image(doc, layers, "图 1 六层架构（Presentation → World）")
    _para(
        doc,
        "自上而下：表现层接收用户输入；控制层定义 ModelPolicy 与模式；编排层产出计划；"
        "执行层是唯一 syscall 出口；状态层持久化记忆；世界层维护纯状态叙事。",
    )

    _heading(doc, "3. 唯一执行路径", 1)
    _add_image(doc, pipeline, "图 2 内核 entry() 流水线")
    _bullet(
        doc,
        [
            "禁止多 ExecutionEngine、多 ToolRouter 主入口",
            "WorldRuntime 禁止直接调 tool / shell / 写文件",
            "Renderer（NarrativeWriter）只读 state 生成文本",
            "显式 memory mutation 仅 MemoryLayer.apply_mutation，且经 ExecutionEngine",
        ],
    )

    _heading(doc, "4. 核心模块映射", 1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "路径"
    hdr[1].text = "职责"
    rows = [
        ("core/runtime/agent_os_runtime.py", "唯一系统入口"),
        ("core/runtime/execution_engine.py", "唯一执行引擎"),
        ("core/memory/memory_layer.py", "唯一记忆入口"),
        ("core/control/model_policy.py", "模型与能力边界"),
        ("core/orchestration/autonomous_loop.py", "Phase5 自主循环"),
        ("agent_service/", "生产 FastAPI"),
        ("web/frontend/", "React 开发者控制台"),
    ]
    for a, b in rows:
        row = table.add_row().cells
        row[0].text = a
        row[1].text = b

    _heading(doc, "5. 参考文档", 1)
    _bullet(doc, ["SYSTEM_BLUEPRINT.md", "ARCHITECTURE_CONSTRAINTS.md", "docs/ARCHITECTURE_DIAGRAM.md（Mermaid）"])

    doc.save(out)
    return out


def build_user_manual_docx(layers: Path) -> Path:
    out = DOCS / "MemoryAgentOS-使用说明.docx"
    doc = Document()
    _heading(doc, "Memory Agent OS — 使用说明", 0)
    _para(doc, "面向演示、测试与录屏的操作手册。")
    doc.add_paragraph()

    _heading(doc, "1. 环境要求", 1)
    _bullet(
        doc,
        [
            "Python 3.10+，推荐 WSL2 Ubuntu 22.04",
            "可选：Ollama + nomic-embed-text（语义记忆）",
            "桌面实况：WSL2 + Windows 主机",
        ],
    )

    _heading(doc, "2. 安装与启动", 1)
    _bullet(
        doc,
        [
            "python3 -m venv .venv && source .venv/bin/activate",
            "pip install -r requirements.txt",
            "./scripts/start_production.sh",
            "浏览器打开 http://127.0.0.1:8787/app/",
        ],
    )

    _heading(doc, "3. 界面操作", 1)
    _para(doc, "录屏请使用开发者模式，以展示完整 trace。")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "输入示例"
    table.rows[0].cells[1].text = "预期"
    for inp, exp in [
        ("1+2", "execute_code，结果 3"),
        ("记住: 我叫张伟强", "execute_memory_op 治理写入"),
        ("叙事类自然语言", "narrative_respond + World 更新"),
        ("Phase5 勾选 + 目标", "多步自主闭环"),
    ]:
        r = table.add_row().cells
        r[0].text = inp
        r[1].text = exp

    _heading(doc, "4. 主要 API", 1)
    _bullet(
        doc,
        [
            "POST /api/run — 单轮执行",
            "POST /api/stream — NDJSON 流式",
            "POST /api/autonomous/run — Phase5 会话",
            "GET /api/desktop/stream — 桌面 MJPEG",
            "POST /api/memory/mutate — 治理化记忆",
            "Swagger: http://127.0.0.1:8787/docs",
        ],
    )

    _heading(doc, "5. CLI", 1)
    _bullet(
        doc,
        [
            "python3 main.py --mode developer --trace",
            'python3 main.py --autonomous-session --autonomous-steps 4 "目标"',
        ],
    )

    _heading(doc, "6. 架构示意", 1)
    _add_image(doc, layers, "图：系统分层", width_in=5.5)

    _heading(doc, "7. 录屏检查清单", 1)
    _bullet(
        doc,
        [
            "pytest tests/ -q 通过",
            "8787 服务可访问",
            "开发者模式已开启",
            "示例输入已准备（见 VIDEO_SCRIPT.md）",
        ],
    )

    doc.save(out)
    return out


def build_test_guide_docx(pipeline: Path) -> Path:
    out = DOCS / "MemoryAgentOS-测试与演示指南.docx"
    doc = Document()
    _heading(doc, "Memory Agent OS — 测试与演示指南", 0)
    doc.add_paragraph()

    _heading(doc, "1. 自动化回归", 1)
    _para(doc, "python3 -m pytest tests/ -q")
    _para(doc, "覆盖：记忆治理、自主循环、Phase4 Guard、ModelPolicy 等。")

    _heading(doc, "2. 冒烟测试", 1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    h = table.rows[0].cells
    h[0].text = "项"
    h[1].text = "命令/操作"
    h[2].text = "通过标准"
    checks = [
        ("服务", "./scripts/start_production.sh", "8787/app 返回 200"),
        ("代码通道", 'POST /api/run {"message":"1+2","mode":"developer"}', "含结果与 trace"),
        ("记忆", 'message: 记住: 测试', "无 500，memory 路由"),
        ("Phase5", "POST /api/autonomous/run steps=2", "多步结构或 trace"),
        ("桌面", "verify_desktop_capture.py", "捕获成功或明确提示"),
    ]
    for a, b, c in checks:
        r = table.add_row().cells
        r[0].text, r[1].text, r[2].text = a, b, c

    _heading(doc, "3. UI 手工测试", 1)
    _bullet(
        doc,
        [
            "开发者模式输入 1+2",
            "输入 记住: 我叫张伟强",
            "勾选 Phase5 执行短目标",
            "（可选）桌面实况刷新",
        ],
    )

    _heading(doc, "4. 推荐演示流程（5 分钟）", 1)
    _bullet(
        doc,
        [
            "启动 start_production.sh",
            "开发者模式：1+2 → 记住 → Phase5",
            "（可选）桌面实况",
            "片尾展示 /docs 或架构 docx",
        ],
    )

    _heading(doc, "5. 录屏脚本", 1)
    _para(doc, "详细分镜见 docs/VIDEO_SCRIPT.md（约 5–8 分钟口播稿）。")

    _heading(doc, "6. 执行路径参考", 1)
    _add_image(doc, pipeline, "图：entry() 流水线")

    doc.save(out)
    return out


def main() -> int:
    DOCS.mkdir(parents=True, exist_ok=True)
    layers_png = DIAGRAMS / "layers.png"
    pipeline_png = DIAGRAMS / "execution_path.png"
    print("Drawing architecture PNGs…")
    draw_layers_diagram(layers_png)
    draw_execution_path(pipeline_png)
    print(f"  {layers_png}")
    print(f"  {pipeline_png}")

    print("Building docx…")
    files = [
        build_architecture_docx(layers_png, pipeline_png),
        build_user_manual_docx(layers_png),
        build_test_guide_docx(pipeline_png),
    ]
    for f in files:
        print(f"  {f}")
    print("Done. Open docs/*.docx for 录屏/答辩交付.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
